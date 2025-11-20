import re
from typing import List
import pdfplumber
from docx import Document
from app.config import settings

def _clean_line(line: str) -> str:
    line = line.replace("\xa0", " ")
    line = re.sub(r'\s+', ' ', line).strip()
    return line

def _is_noise_line(line: str) -> bool:
    if not line:
        return True
    if re.match(r'^\s*page\s*\d+\s*$', line, re.I):
        return True
    if re.match(r'^\s*\d+\s*$', line):
        return True
    non_alnum = len(re.sub(r'[A-Za-z0-9]', '', line))
    if non_alnum / max(len(line), 1) > 0.5:
        return True
    if re.search(r'(table of contents|contents|index|references)$', line.lower()):
        return True
    return False

def _remove_repeated_headers(page_lines_list: List[List[str]]) -> List[str]:
    header_counts = {}
    footer_counts = {}
    pages = len(page_lines_list)
    for lines in page_lines_list:
        if not lines:
            continue
        top = lines[0][:120].strip()
        bottom = lines[-1][:120].strip()
        if top:
            header_counts[top] = header_counts.get(top, 0) + 1
        if bottom:
            footer_counts[bottom] = footer_counts.get(bottom, 0) + 1
    headers_to_remove = {h for h, c in header_counts.items() if c > max(1, pages//2)}
    footers_to_remove = {f for f, c in footer_counts.items() if c > max(1, pages//2)}
    out = []
    for lines in page_lines_list:
        for i, l in enumerate(lines):
            if i == 0 and l in headers_to_remove:
                continue
            if i == len(lines)-1 and l in footers_to_remove:
                continue
            out.append(l)
    return out

def extract_text_from_pdf(path: str) -> str:
    page_lines = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            raw = page.extract_text() or ""
            lines = [_clean_line(ln) for ln in raw.splitlines()]
            lines = [ln for ln in lines if ln]  # drop empty
            page_lines.append(lines)
    lines = _remove_repeated_headers(page_lines)
    lines = [l for l in lines if not _is_noise_line(l)]
    lines = [re.sub(r'\s{2,}', ' ', l) for l in lines]
    text = "\n".join(lines).strip()
    text = re.sub(r'\n{2,}', '\n\n', text)
    return text

def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    paras = []
    for p in doc.paragraphs:
        t = _clean_line(p.text)
        if t and not _is_noise_line(t):
            paras.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [ _clean_line(cell.text) for cell in row.cells if _clean_line(cell.text) ]
            if cells:
                paras.append(" | ".join(cells))
    text = "\n".join(paras)
    text = re.sub(r'\n{2,}', '\n\n', text)
    return text

def extract_text_from_file(path: str) -> str:
    lower = path.lower()
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(path)
    if lower.endswith(".docx") or lower.endswith(".doc"):
        return extract_text_from_docx(path)
    return extract_text_from_pdf(path)
